import csv
from docx import Document
from docx.shared import Cm
import os
import win32com.client
from commonUtils import configParams

config = configParams()
params = config.loadConfig()

def read_csv_to_list(file_path):
    data_list = []
    
    with open(file_path, 'r') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            data_list.append(row)
    
    return data_list

def update_word_document(template_path, output_path, data):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in ["Porta", "Pis"] and not value:
                value = "--"
            if "{{" + key + "}}" in paragraph.text:
                if key == "Imatge" and value:
                    # Assuming the placeholder for image is a complete line on its own
                    paragraph.text = paragraph.text.replace("{{" + key + "}}", "")     
                    run = paragraph.add_run()
                    sign = params["carpetaSignatures"] + f"\{value}"
                    run.add_picture(sign, Cm(10))  # Adjust width as required
                    
                else:
                    paragraph.text = paragraph.text.replace("{{" + key + "}}", value)

    doc.save(output_path)
    #print(f"DOCX Generat: {output_path}")

def saveDocX_as_PDF(input_file, output_file):
    current_script_directory = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(current_script_directory, input_file)
    output_path = os.path.join(current_script_directory, output_file)

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)  # 17 represents the PDF format in Word
    doc.Close()
    word.Quit()
    print(f"PDF Generat: {output_path}")

peticions = read_csv_to_list(params["csvPeticions"])
signatures = read_csv_to_list(params["csvSignatures"])

for row in signatures:
    for psf in peticions:
        if row['PDF'] != "Y":
            codi = row['Codi']
            document = psf["Document"]
            baseDOCX = f'{params["carpetaDocumentsPeticions"]}\{document}'
            wordDoc = f'{params["carpetaDOCXs"]}\{codi}-{document}'
            update_word_document(baseDOCX, wordDoc, row)
            pdfDoc = f'{params["carpetaPDFs"]}\{os.path.splitext(os.path.basename(wordDoc))[0]}.pdf'
            saveDocX_as_PDF(wordDoc, pdfDoc)


